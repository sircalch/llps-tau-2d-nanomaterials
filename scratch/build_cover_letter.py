import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def build_cover_letter():
    md_content = """# Cover Letter for Soft Matter (Royal Society of Chemistry)

**Date:** September 3, 2026  
**To:** The Editorial Office, *Soft Matter*  
**From:** Prof. Andrés Monreal Hernández (Corresponding Author)  
Universidad Estatal de Sonora, Hermosillo, Sonora, Mexico  
Email: andres.monreal@ues.mx  

**Subject:** Submission of Original Research Article:  
*"Thermodynamic modulation of Tau liquid-liquid phase separation and condensate wetting by two-dimensional nanomaterial interfaces: emergent suppression via adsorption equilibrium"*

Dear Editor,

We are pleased to submit our manuscript entitled **"Thermodynamic modulation of Tau liquid-liquid phase separation and condensate wetting by two-dimensional nanomaterial interfaces: emergent suppression via adsorption equilibrium"** for consideration as an Original Research Paper in *Soft Matter*.

### Motivation and Core Soft-Matter Insight
Biomolecular condensates formed by liquid-liquid phase separation (LLPS) of intrinsically disordered proteins (IDPs) such as Tau regulate cellular compartmentalization, but aberrant phase transitions inside crowded droplets can nucleate pathological cross-beta amyloid fibrils associated with tauopathies. Despite intense interest in soft nanotechnology and biointerfaces, a predictive theoretical framework explaining how structured physical boundaries modulate biomolecular phase coexistence, wetting transitions, and condensate aging kinetics has remained lacking.

In this work, we present a coarse-grained statistical-thermodynamic and kinetic theory that couples Flory-Huggins-Voorn-Overbeek (FH-VO) polymer physics with Langmuir interfacial adsorption mass balance, Cahn-Hilliard wetting theory, and strictly dimensional master equations. 

The central soft-matter insight of our study is that **2D interface-mediated LLPS suppression emerges self-consistently from interfacial monomer sequestration governed by the interfacial area density (a_s), without requiring empirical alterations to the intrinsic macromolecular Flory interaction parameter (dchi/da_s = 0)**.

### Highlights of the Framework:
1. **Calibrated to Experimental Tau K18 Thermodynamics:** The bulk phase model is quantitatively calibrated against published Lower Critical Solution Temperature (LCST) turbidity onset measurements of the microtubule-binding repeat domain Tau K18 (15.3 deg C at 100 uM; Ambadipudi et al., Nat. Commun. 2017).
2. **General Interfacial Coordinate & Representative Scenarios:** We establish the interfacial area density a_s as the fundamental physical coordinate. We evaluate two literature-informed representative contrast scenarios: a high-affinity borophene-like scenario (dG_ads = -7.8 kcal/mol, contact angle theta_c = 50.3 deg) that dissolves LLPS up to 29.4 deg C at a_s = 1.0e-4 nm^-1 and depletes ~60% of free monomers at 37 deg C; and a moderate-affinity Ti3C2Tx MXene-like scenario (dG_ads = -5.2 kcal/mol, theta_c = 79.3 deg) that maintains droplet coexistence.
3. **Dimensional Aging Kinetics:** Holding the underlying aggregation rate-law structure fixed, dimensional master equations demonstrate that interfacial monomer extraction selectively retards autocatalytic secondary nucleation inside condensates.
4. **Variance-Based Global Sensitivity Analysis (SALib):** A full Saltelli Sobol analysis (N_base = 512, N_eval = 5120) with nested block convergence reveals a clear mechanistic partition: thermal slope beta, ionic strength I, and segment length N_eff form an overlapping thermodynamic cluster governing phase coexistence, whereas interfacial area density a_s and extraction kinetics k_ext dictate fibrillation arrest.

### Compliance with RSC Policies and Declarations:
- **Scope Fit:** This manuscript fits squarely within the scope of *Soft Matter*, bridging polymer theory, complex coacervation, interfacial wetting, and soft nanotechnology.
- **Conflicts of Interest:** The authors declare that there are no conflicts of interest.
- **Open Data & Reproducibility:** All simulation codes, root solvers, kinetic master equation integrators, digitized data, and automated test suites are openly accessible in the project repository: https://github.com/sircalch/llps-tau-2d-nanomaterials. An immutable Zenodo DOI will be permanently minted from the tagged release candidate (v1.0.0).
- **Generative AI Disclosure:** In strict accordance with RSC author guidelines, generative AI tools (OpenAI Codex / Anthropic Claude / Google Gemini) were utilized during computational development for code review, unit test generation, and manuscript formatting. All mathematical models, numerical simulations, physical interpretations, and citations were independently formulated, verified, and approved by the authors.

### Suggested Independent Referees:
1. **Prof. Rohit V. Pappu** (Washington University in St. Louis) - Expert in intrinsically disordered proteins, polymer physics of biomolecular condensates, and phase behavior. Email: pappu@wustl.edu
2. **Prof. Tuomas P. J. Knowles** (University of Cambridge) - Expert in protein aggregation kinetics, biomolecular phase transitions, and amyloid nucleation theory. Email: tpjk2@cam.ac.uk
3. **Prof. Yury Gogotsi** (Drexel University) - Expert in 2D nanomaterials, MXene surface chemistry, and nanomaterial-biomolecule interfaces. Email: gogotsi@drexel.edu
4. **Dr. Roland L. Knorr** (Max Planck Institute of Colloids and Interfaces) - Expert in biomolecular condensate wetting, interfacial tension, and membrane interactions. Email: roland.knorr@mpikg.mpg.de

Thank you very much for your time and editorial consideration.

Sincerely,

**Prof. Andrés Monreal Hernández**  
Corresponding Author  
Universidad Estatal de Sonora  
Hermosillo, Sonora, Mexico  
Email: andres.monreal@ues.mx
"""
    os.makedirs("manuscript", exist_ok=True)
    with open("manuscript/cover_letter_Soft_Matter.md", "w", encoding="utf-8") as f:
        f.write(md_content)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

    p = doc.add_paragraph()
    r = p.add_run("COVER LETTER")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_after = Pt(12)

    def add_p(text, bold=False, italic=False, space_after=6, line_spacing=1.15):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.bold = bold; r.font.italic = italic
        return p

    add_p("Date: September 3, 2026", bold=True)
    add_p("To: The Editorial Office, Soft Matter (Royal Society of Chemistry)", bold=True)
    add_p("From: Prof. Andrés Monreal Hernández (Corresponding Author)\nUniversidad Estatal de Sonora, Hermosillo, Sonora, Mexico\nEmail: andres.monreal@ues.mx")
    add_p("Subject: Submission of Original Research Article", bold=True, space_after=12)

    add_p("Dear Editor,", space_after=8)
    add_p("We are pleased to submit our manuscript entitled \"Thermodynamic modulation of Tau liquid-liquid phase separation and condensate wetting by two-dimensional nanomaterial interfaces: emergent suppression via adsorption equilibrium\" for consideration as an Original Research Paper in Soft Matter.")

    add_p("Motivation and Core Soft-Matter Insight:", bold=True, space_after=4)
    add_p("Biomolecular condensates formed by liquid-liquid phase separation (LLPS) of intrinsically disordered proteins (IDPs) such as Tau regulate cellular compartmentalization, but aberrant phase transitions inside crowded droplets can nucleate pathological cross-beta amyloid fibrils associated with tauopathies. Despite intense interest in soft nanotechnology and biointerfaces, a predictive theoretical framework explaining how structured physical boundaries modulate biomolecular phase coexistence, wetting transitions, and condensate aging kinetics has remained lacking.")
    add_p("In this work, we present a coarse-grained statistical-thermodynamic and kinetic theory coupling Flory-Huggins-Voorn-Overbeek (FH-VO) polymer physics with Langmuir interfacial adsorption mass balance, Cahn-Hilliard wetting theory, and strictly dimensional master equations. The central soft-matter insight of our study is that 2D interface-mediated LLPS suppression emerges self-consistently from interfacial monomer sequestration governed by the interfacial area density (a_s), without requiring empirical alterations to the intrinsic macromolecular Flory interaction parameter (dchi/da_s = 0).")

    add_p("Highlights of the Framework:", bold=True, space_after=4)
    add_p("1. Quantitative Calibration to Tau K18: Parameterized against experimental LCST turbidity onset measurements of Tau K18 (15.3 °C at 100 µM; Ambadipudi et al., Nat. Commun. 2017).")
    add_p("2. Invariant Interfacial Coordinate & Representative Scenarios: Establishes interfacial area density a_s as the fundamental physical coordinate. We contrast a high-affinity borophene-like scenario (dG_ads = -7.8 kcal/mol, contact angle theta_c = 50.3°) that dissolves LLPS up to 29.4 °C at a_s = 1.0e-4 nm^-1 and depletes ~60% monomer at 37 °C, against a moderate-affinity Ti3C2Tx MXene-like scenario (dG_ads = -5.2 kcal/mol, theta_c = 79.3°) that maintains droplet coexistence.")
    add_p("3. Dimensional Aging Kinetics: Demonstrates that interfacial monomer extraction delays autocatalytic secondary nucleation inside condensates without altering intrinsic rate laws.")
    add_p("4. Saltelli Sobol Global Sensitivity Analysis (SALib): With N_base = 512 (5120 physical model evaluations) and nested block convergence, showing that thermal slope beta, ionic strength I, and segment length N_eff form an overlapping thermodynamic cluster controlling phase coexistence, while area density a_s and extraction kinetics k_ext govern fibrillation arrest.")

    add_p("Compliance with RSC Policies and Declarations:", bold=True, space_after=4)
    add_p("• Scope Fit: Directly aligns with the scope of Soft Matter, bridging biopolymers, liquid-liquid phase coexistence, wetting transitions, and soft nanotechnology.")
    add_p("• Conflicts of Interest: The authors declare that there are no conflicts of interest.")
    add_p("• Open Data & Code: Fully open in GitHub (https://github.com/sircalch/llps-tau-2d-nanomaterials) with continuous integration. An immutable Zenodo DOI will be permanently minted from the v1.0.0 release.")
    add_p("• Generative AI Disclosure: In accordance with RSC policy, generative AI tools (OpenAI Codex / Anthropic Claude / Google Gemini) were utilized for code review, unit test generation, and manuscript formatting. All models, calculations, and interpretations were independently formulated and verified by the authors.")

    add_p("Suggested Independent Referees:", bold=True, space_after=4)
    add_p("1. Prof. Rohit V. Pappu (Washington University in St. Louis) - pappu@wustl.edu\n2. Prof. Tuomas P. J. Knowles (University of Cambridge) - tpjk2@cam.ac.uk\n3. Prof. Yury Gogotsi (Drexel University) - gogotsi@drexel.edu\n4. Dr. Roland L. Knorr (Max Planck Institute of Colloids and Interfaces) - roland.knorr@mpikg.mpg.de")

    add_p("Thank you very much for your time and editorial consideration.\n\nSincerely,\nProf. Andrés Monreal Hernández\nCorresponding Author", space_after=12)

    doc_out = "manuscript/cover_letter_Soft_Matter.docx"
    doc.save(doc_out)
    print(f"Cover letter created: {doc_out} and manuscript/cover_letter_Soft_Matter.md")

if __name__ == "__main__":
    build_cover_letter()
