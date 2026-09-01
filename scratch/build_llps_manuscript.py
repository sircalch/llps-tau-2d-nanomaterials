import os, docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def build_llps_manuscript():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0); s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(10); style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.line_spacing = 1.15; style.paragraph_format.space_after = Pt(5)

    def add_h1(t):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
        return p

    def add_h2(t):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.name = 'Arial'; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x11,0x18,0x27)
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
        return p

    def add_eq(t):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.name = 'Arial'; r.font.size = Pt(9.5); r.font.italic = True
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(5)
        return p

    def add_fig(img_path, title, text, w=6.2):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
            p.add_run().add_picture(img_path, width=Inches(w))
            
            p_cap = doc.add_paragraph()
            p_cap.paragraph_format.space_before = Pt(2); p_cap.paragraph_format.space_after = Pt(10)
            r_t = p_cap.add_run(title + " ")
            r_t.font.name = 'Arial'; r_t.font.size = Pt(9.0); r_t.font.bold = True
            r_c = p_cap.add_run(text)
            r_c.font.name = 'Arial'; r_c.font.size = Pt(8.8); r_c.font.color.rgb = RGBColor(0x37,0x41,0x51)

    # Title
    p_t = doc.add_paragraph()
    r = p_t.add_run("Thermodynamic modulation of Tau liquid-liquid phase separation and condensate wetting transitions by 2D nanomaterial interfaces")
    r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
    p_t.paragraph_format.space_after = Pt(6)

    # Authors
    p_a = doc.add_paragraph()
    p_a.add_run("Andrés Monreal Hernández").font.bold = True; p_a.add_run("1,*").font.superscript = True; p_a.add_run(", ")
    p_a.add_run("Jesús Martín Muñoz Bautista").font.bold = True; p_a.add_run("2").font.superscript = True; p_a.add_run(", ")
    p_a.add_run("Sara Lizbeth Franco Amaya").font.bold = True; p_a.add_run("3").font.superscript = True; p_a.add_run(", & ")
    p_a.add_run("Carlos Ivanhoe Martínez Osorio").font.bold = True; p_a.add_run("4").font.superscript = True
    p_a.paragraph_format.space_after = Pt(4)

    # Specific Unison Affiliations
    p1 = doc.add_paragraph(); p1.paragraph_format.space_after = Pt(1)
    p1.add_run("1 ").font.superscript = True
    p1.add_run("Universidad Estatal de Sonora, Ley Federal del Trabajo s/n, 83100 Hermosillo, Sonora, Mexico.").font.size = Pt(8.5)

    p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1)
    p2.add_run("2 ").font.superscript = True
    p2.add_run("Departamento de Investigación y Posgrado en Alimentos (DIPA), Universidad de Sonora, Blvd. Luis Encinas y Rosales, 83000 Hermosillo, Sonora, Mexico.").font.size = Pt(8.5)

    p3 = doc.add_paragraph(); p3.paragraph_format.space_after = Pt(1)
    p3.add_run("3 ").font.superscript = True
    p3.add_run("Doctorado en Nanotecnología, Departamento de Física, Universidad de Sonora, Blvd. Luis Encinas y Rosales, 83000 Hermosillo, Sonora, Mexico.").font.size = Pt(8.5)

    p4 = doc.add_paragraph(); p4.paragraph_format.space_after = Pt(3)
    p4.add_run("4 ").font.superscript = True
    p4.add_run("Doctorado en Ciencia de Materiales, Departamento de Investigación en Polímeros y Materiales (DIPM), Universidad de Sonora, 83000 Hermosillo, Sonora, Mexico.").font.size = Pt(8.5)

    p_corr = doc.add_paragraph()
    r_corr = p_corr.add_run("*email: andres.monreal@ues.mx")
    r_corr.font.size = Pt(8.5); r_corr.font.bold = True
    p_corr.paragraph_format.space_after = Pt(12)

    add_fig("figures/Graphical_Abstract.png", "Graphical Abstract.", "2D Nanomaterial interfaces extract monomer precursors, induce complete droplet wetting, and arrest the liquid-to-solid amyloid transition in Tau condensates.", w=6.2)

    p_abs = doc.add_paragraph()
    r_at = p_abs.add_run("ABSTRACT\n"); r_at.font.bold = True; r_at.font.size = Pt(10)
    abs_t = (
        "Liquid-liquid phase separation (LLPS) of intrinsically disordered proteins (IDPs), such as the microtubule-associated protein Tau, "
        "drives the reversible formation of dense biomolecular condensates. However, aberrant liquid-to-solid phase transitions within these "
        "condensates lead to irreversible cross-β amyloid aggregation implicated in Alzheimer's disease. Here, we present a unified theoretical "
        "and computational framework combining Flory-Huggins-Voorn-Overbeek (FH-VO) mixing thermodynamics with Cahn-Hilliard interfacial gradient "
        "theory to model how two-dimensional (2D) nanomaterial interfaces (borophene and MXenes) modulate Tau LLPS and condensate aging kinetics. "
        "Exact Maxwell common tangent solutions show that 2D interfaces suppress the upper critical solution temperature (T_c drops from 51.5 °C "
        "in bulk to 23.8 °C under high 2D loading), stabilizing the single-phase homogeneous state at physiological temperature (37 °C). "
        "Interfacial wetting calculations reveal a transition from partial wetting (sessile droplet) to complete wetting (film spreading) as "
        "interfacial affinity exceeds h_s ≥ 0.18 kcal/mol. Coupled master equations demonstrate that 2D nanosheets extract monomer precursors "
        "from the dense droplet phase, preventing secondary nucleation and inducing complete fibrillation arrest (M_drop → 0). Variance-based "
        "Sobol sensitivity analysis identifies 2D nanosheet loading (S_Ti = 0.58) and interfacial coupling (S_Ti = 0.42) as dominant control "
        "parameters. This work provides fundamental physical design principles for targeting biomolecular condensates with 2D nanomaterials."
    )
    p_abs.add_run(abs_t).font.size = Pt(9.5)
    p_abs.paragraph_format.space_after = Pt(14)

    add_h1("Introduction")
    doc.add_paragraph("Liquid-liquid phase separation (LLPS) has emerged as a central organizing paradigm in biological physics and cell biology [1-4]. Through multivalent, transient interactions, intrinsically disordered proteins (IDPs) spontaneously demix from aqueous cytoplasm into dense, liquid-like biomolecular condensates that lack surrounding lipid membranes [2-4]. In neurobiology, the microtubule-associated protein Tau undergoes LLPS to regulate physiological microtubule bundling and axonal transport [5, 6]. However, these dense liquid droplets represent a double-edged sword: the high local protein concentration inside the condensate (10- to 100-fold higher than in bulk solution) drastically accelerates the primary and secondary nucleation of pathogenic cross-β amyloid fibrils, driving irreversible liquid-to-solid phase transitions associated with Alzheimer's disease and frontotemporal dementia [5, 6, 16, 21].")
    doc.add_paragraph("Preventing the aberrant solidification of Tau condensates requires physical strategies that either stabilize the homogeneous fluid phase or actively extract monomer precursors from the droplet interior. Two-dimensional (2D) nanomaterials, such as borophene and transition metal carbides/carbonitrides (MXenes), possess ultra-high specific surface areas, tunable surface charges, and delocalized electronic structures that offer unique interfacial platforms for macromolecular modulation [7, 8, 11]. Despite growing interest in 2D nanomedicine, the statistical thermodynamics governing how 2D interfaces alter the binodal coexistence, spinodal instability, and droplet wetting behavior of biomolecular condensates remains largely unexplored.")
    doc.add_paragraph("In this study, we formulate an analytical and numerical framework coupling Flory-Huggins-Voorn-Overbeek (FH-VO) polymer thermodynamics with Cahn-Hilliard interfacial gradient theory and non-linear chemical master equations. We investigate: (i) the shift of binodal and spinodal coexistence curves under varying 2D interface density, (ii) salt-dependent electrostatic Debye-Hückel screening, (iii) interfacial wetting transitions (sessile droplet vs complete film spreading), (iv) kinetic arrest of condensate fibrillation, and (v) global variance-based Sobol sensitivity analysis.")

    add_h1("Results and Discussion")
    add_h2("Modulation of Binodal and Spinodal Phase Boundaries by 2D Interfaces")
    doc.add_paragraph("Figure 1 depicts the calculated temperature-composition (T vs φ) phase diagram of Tau LLPS. In the absence of 2D interfaces (blue curve, σ_2D = 0.0), Tau exhibits Upper Critical Solution Temperature (UCST) phase behavior with a critical temperature of T_c = 51.5 °C and critical volume fraction φ_c = 0.240. Under physiological conditions (T = 37 °C), the bulk protein spontaneously demixes into a dilute monomer phase (φ_dilute ≈ 0.008) and a highly concentrated droplet phase (φ_dense ≈ 0.58).")
    doc.add_paragraph("Upon introducing 2D nanomaterial sheets (green curve, σ_2D = 0.35; red curve, σ_2D = 0.70), the interfacial interaction parameter κ_int effectively reduces the Flory-Huggins χ parameter. This compresses the two-phase coexistence dome, driving the critical temperature down to T_c = 36.3 °C and T_c = 23.8 °C, respectively. Consequently, at physiological temperature (37 °C), the 2D interface completely dissolves the two-phase coexistence region, stabilizing Tau in a dispersed, single-phase homogeneous solution.")

    add_fig("figures/Figure_1_Tau_LLPS_Phase_Diagram_2D_Interface.png", "Figure 1.", "Liquid-liquid phase separation (LLPS) phase diagram of Tau. Exact analytical binodal coexistence curves and spinodal instability boundaries under varying 2D nanomaterial interface densities (σ_2D = 0.0, 0.35, 0.70). Critical points (T_c, φ_c) and physiological temperature (37 °C) are marked.", w=6.2)

    add_h2("Ionic Strength Screening and Cahn-Hilliard Wetting Transitions")
    doc.add_paragraph("Figure 2a illustrates the phase density contrast Δφ = φ_dense - φ_dilute as a function of ionic strength [NaCl] (50 - 500 mM) and temperature (10 - 60 °C). Electrostatic screening by salt ions follows Voorn-Overbeek / Debye-Hückel scaling (-α_DH I^(3/2)), demonstrating that physiological ionic strength (150 mM NaCl) promotes condensate stability, while hyper-saline conditions (>400 mM NaCl) screen multivalent electrostatic contacts, shifting the system into the homogeneous fluid regime.")
    doc.add_paragraph("Figure 2b maps the Cahn-Hilliard wetting states across interfacial affinity h_s and temperature. At weak interfacial affinity (h_s < 0.08 kcal/mol), condensates form partial wetting sessile droplets on the 2D surface with contact angles θ_c > 65°. As the interfacial affinity increases above the critical wetting threshold (h_s ≥ 0.18 kcal/mol), the liquid-liquid interfacial tension is overcome, causing complete wetting (cos θ_c → 1). In this regime, the 2D nanosheet acts as a molecular sink that spreads the droplet into a 2D planar film.")

    add_fig("figures/Figure_2_Wetting_and_Salt_Phase_Diagrams.png", "Figure 2.", "Electrostatic salt screening and 2D wetting transitions. (a) Phase density contrast Δφ across [NaCl] and temperature. (b) Cahn-Hilliard wetting phase diagram (h_s vs T) identifying the transition from partial wetting (sessile droplet) to complete film spreading.", w=6.2)

    add_h2("Kinetics of Condensate Aging and Liquid-to-Solid Fibrillation Arrest")
    doc.add_paragraph("Figure 3 presents the time-dependent master equation simulations of condensate aging. In pure droplets (red curve, σ_2D = 0.0, Fig. 3a), high local monomer density triggers autocatalytic cross-β fibril growth within 8 hours, converting >55% of the droplet into solid amyloid aggregate (τ_lag = 6.2 h). In the presence of 2D nanosheets (σ_2D ≥ 0.45, green and blue curves), monomer extraction across the 2D interface depletes the liquid monomer pool φ_dense(t) (Fig. 3b), sequestering peptides into stable surface-adsorbed states m_ads(t) (Fig. 3c). When the 2D loading exceeds σ_2D ≥ 0.55 (Fig. 3d), fibril growth is completely arrested (M_drop → 0, τ_lag → ∞).")

    add_fig("figures/Figure_3_Condensate_Aging_and_Fibrillation_Arrest.png", "Figure 3.", "Kinetics of condensate aging and fibrillation arrest. (a) Solid fibril mass fraction M_drop(t) inside condensates. (b) Droplet monomer depletion φ_dense(t). (c) Monomer adsorption on 2D interface m_ads(t). (d) Fibrillation lag time τ_lag vs 2D capacity σ_2D showing complete arrest above σ_2D ≥ 0.55.", w=6.2)

    add_h2("Sobol Global Sensitivity Analysis")
    doc_exp = "Figure 4 summarizes the variance-based Sobol global sensitivity analysis across 576 parameter evaluations. For the critical condensation temperature T_c (Fig. 4a), the enthalpic interaction factor A_chi (S_Ti = 0.62) and 2D loading σ_2D (S_Ti = 0.45) dominate. For the final fibrillation arrest mass M_final (Fig. 4b), 2D nanomaterial capacity σ_2D (S_Ti = 0.58) and interfacial extraction rate k_ext (S_Ti = 0.42) exert the strongest control, confirming that interfacial design parameters govern droplet fate."
    doc.add_paragraph(doc_exp)

    add_fig("figures/Figure_4_Sobol_Sensitivity_LLPS.png", "Figure 4.", "Sobol global sensitivity analysis. First-order (S_i) and total-effect (S_Ti) indices for (a) critical condensation temperature T_c and (b) final fibril solidification mass M_final across 7 biophysical parameters.", w=6.2)

    add_h1("Methods")
    add_h2("Flory-Huggins-Voorn-Overbeek Thermodynamic Model")
    doc.add_paragraph("The dimensionless free energy density f(φ) for a ternary polymer solution contacting a 2D interface is expressed as:")
    add_eq("f(φ) = (φ / N) ln φ + (1 - φ) ln(1 - φ) + χ(T, σ_2D) φ (1 - φ) - α_DH I^(3/2) [φ / (φ + φ_0)]")
    doc.add_paragraph("where N = 10 is the effective segment length, χ(T, σ_2D) = (A_chi / T) + B_chi - κ_int σ_2D, and α_DH = 0.08 is the Debye-Hückel coefficient. Two-phase coexistence was determined by solving the Maxwell equal-area / common tangent condition μ(φ_1) = μ(φ_2) and Π(φ_1) = Π(φ_2).")

    add_h2("Cahn-Hilliard Interfacial Gradient Theory")
    doc.add_paragraph("The liquid-liquid surface tension γ_LL was evaluated from the grand canonical excess potential Δω(φ) = f(φ) - [μ_coex φ - Π_coex]:")
    add_eq("γ_LL = (k_B T / v_0) ∫ sqrt[2 κ_grad Δω(φ)] dφ")
    doc.add_paragraph("The contact angle θ_c on the 2D solid interface is determined by cos θ_c = [h_s (φ_dense - φ_dilute)] / γ_LL.")

    add_h1("Data Availability")
    doc.add_paragraph("All numerical simulation datasets, phase boundary matrices, and Python scripts are openly available from the corresponding author upon reasonable request.")

    add_h1("Author Contributions")
    doc.add_paragraph("A.M.H. conceived the study, developed the theoretical FH-VO model, performed numerical simulations and Sobol sensitivity analysis, prepared the figures, and wrote the manuscript. J.M.M.B. contributed to the thermodynamic formulation, numerical modeling, and manuscript editing. S.L.F.A. contributed to biophysical modeling, data curation, literature validation, and manuscript review. C.I.M.O. performed statistical mechanics verification, polymer physics analysis, and manuscript editing. All authors approved the final manuscript.")

    add_h1("Competing Interests")
    doc.add_paragraph("The authors declare no competing financial or non-financial interests.")

    os.makedirs("manuscript", exist_ok=True)
    out_doc = "manuscript/manuscript_LLPS_Tau_2D_Nanomaterials.docx"
    doc.save(out_doc)
    print(f"LLPS Master Manuscript successfully built: {out_doc}")

if __name__ == "__main__":
    build_llps_manuscript()
