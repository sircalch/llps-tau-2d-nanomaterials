"""
build_submission_package.py
===========================
Assembles submission/upload/ with the exact set of files to upload to the
Royal Society of Chemistry submission system (ScholarOne) for a Soft Matter Paper,
using RSC-friendly filenames, plus the editable graphical-abstract text .docx.

Run after the figures and manuscript have been rebuilt (run_pipeline.py does this).
"""

import os, shutil
from docx import Document
from docx.shared import Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUT = os.path.join(ROOT, "submission", "upload")

# RSC accepts vector PDF for figures (TIFF 600 dpi is the alternative); the
# matplotlib figures are vector, so the PDFs are the primary upload artefacts.
FIGURE_MAP = [
    ("figures/Figure_1_Tau_LLPS_Phase_Diagram.pdf",        "Figure1.pdf"),
    ("figures/Figure_2_Wetting_and_Salt_Phase_Diagrams.pdf","Figure2.pdf"),
    ("figures/Figure_3_Borophene_vs_MXene_Comparison.pdf",  "Figure3.pdf"),
    ("figures/Figure_4_Condensate_Aging_Kinetics.pdf",      "Figure4.pdf"),
    ("figures/Figure_5_Sobol_Sensitivity_Analysis.pdf",     "Figure5.pdf"),
]
OTHER_MAP = [
    ("manuscript/manuscript_LLPS_Tau_2D_Nanomaterials.docx", "Manuscript.docx"),
    ("manuscript/cover_letter_Soft_Matter.docx",             "CoverLetter.docx"),
    ("figures/TOC_Graphic_RSC_Soft_Matter.tif",              "GraphicalAbstract.tif"),
    ("figures/TOC_Graphic_RSC_Soft_Matter.pdf",              "GraphicalAbstract.pdf"),
]

# RSC graphical-abstract text: 1-2 sentences, <= 250 characters, focused on the
# key finding and its importance, NOT a caption and NOT a paraphrase of the title/abstract.
GA_TEXT = (
    "High-surface-area 2D nanosheets suppress Tau liquid-liquid phase separation purely "
    "by adsorbing free protein at their interface; the same interfacial sequestration "
    "also stalls the condensates' liquid-to-amyloid ageing."
)


def build_ga_docx(path):
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Graphical abstract text (Soft Matter, <=250 characters)")
    r.bold = True; r.font.size = Pt(11)
    doc.add_paragraph(GA_TEXT)
    n = doc.add_paragraph()
    rn = n.add_run(f"[character count: {len(GA_TEXT)}]")
    rn.italic = True; rn.font.size = Pt(9)
    doc.save(path)


def main():
    if os.path.isdir(OUT):
        shutil.rmtree(OUT)
    os.makedirs(OUT)

    manifest = []
    for src, dst in FIGURE_MAP + OTHER_MAP:
        s = os.path.join(ROOT, src)
        if not os.path.exists(s):
            print(f"  WARNING missing: {src}")
            continue
        shutil.copy2(s, os.path.join(OUT, dst))
        manifest.append((dst, os.path.getsize(s)))

    ga_path = os.path.join(OUT, "GraphicalAbstract_text.docx")
    build_ga_docx(ga_path)
    manifest.append(("GraphicalAbstract_text.docx", os.path.getsize(ga_path)))

    assert len(GA_TEXT) <= 250, f"GA text is {len(GA_TEXT)} chars (limit 250)"

    print(f"Submission upload bundle assembled in {os.path.relpath(OUT, ROOT)}/")
    for name, size in sorted(manifest):
        print(f"  {name:32s} {size/1024:8.1f} KB")
    print(f"\nGraphical-abstract text: {len(GA_TEXT)} / 250 characters")


if __name__ == "__main__":
    main()
